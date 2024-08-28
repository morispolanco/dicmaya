import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO

# Set page configuration
st.set_page_config(page_title="Diccionario Cultural Maya", page_icon="📚", layout="wide")

# API keys
TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
SERPLY_API_KEY = st.secrets["SERPLY_API_KEY"]

# 101 cultural terms related to the Maya perspective, ordered alphabetically
terminos_culturales = sorted([
    "Agua", "Ajaw", "Alma", "Animales", "Árboles", "Balam (Jaguar)", "Ceiba (Árbol sagrado)",
    "Chaac (Dios de la lluvia)", "Ch'ulel (Espíritu o alma)", "Comunión", "Confianza", "Cosmovisión", "Creador",
    "Creación", "Crecimiento", "Cultivo", "Destino", "Dualidad", "Dzuli (Extranjero)", "Etnia", "Espíritu",
    "Espacio-tiempo", "Familia", "Fertilidad", "Fuego", "Fuerza", "Hunab Ku (Dios supremo)", "Hombres",
    "Identidad", "Ik' (Viento)", "Inframundo", "Itzamná (Dios del cielo)", "Ixchel (Diosa de la luna)",
    "Jaguar", "Ja' (Agua)", "K'uh (Divinidad)", "Kukulkan", "Kukulcán (Serpiente emplumada)", "Lak'in (Punto cardinal este)",
    "Luna", "Madera", "Maíz (Elemento vital)", "Maya", "Muerte", "Nahualismo (Creencia en el espíritu animal)",
    "Nación", "Naturaleza", "Noche", "Noj (Sabiduría)", "Origen", "Oxlahun ti' k'uh (Trece dioses)",
    "Pawahtun (Dioses del inframundo)", "Paz", "Persona", "Poder", "Pueblo", "Puntos cardinales", "Q'anil (Semilla)",
    "Religión", "Reencarnación", "Respiración", "Ritual", "Sabiduría", "Sak Nikte' (Flor blanca)",
    "Sak'ij (Iluminación)", "Salud", "Serpiente emplumada", "Sol", "Suerte", "Sueño", "Tezcatlipoca (Espejo humeante)",
    "Tierra", "Tikal (Ciudad maya)", "Tiempo", "Toj (Tributo)", "Tradición", "Trascendencia", "Tzolkin (Calendario sagrado)",
    "Uinal (Mes maya)", "Uxmal (Tres veces construido)", "Universo", "Viento", "Vida", "Visión", "Vitalidad",
    "Voluntad", "Wak (Seis)", "Xenil (Camino)", "Xibalbá (Inframundo maya)", "Yaxché (Árbol de la vida)",
    "Yuum Kaax (Dios de la flora y fauna)", "Yuum k’aax (Dios de los bosques)", "Yuum ik’ (Dios del viento)", 
    "Zama (Amanecer)"
])

def buscar_informacion(query):
    try:
        url = f"https://api.serply.io/v1/scholar/q={query}"
        headers = {
            'X-Api-Key': SERPLY_API_KEY,
            'Content-Type': 'application/json',
            'X-Proxy-Location': 'US',
            'X-User-Agent': 'Mozilla/5.0'
        }
        response = requests.get(url, headers=headers)
        response.raise_for_status()  # Check for HTTP errors
        return response.json()  # Attempt to parse JSON
    except requests.exceptions.RequestException as e:
        st.error(f"Error de red al buscar información para '{query}': {e}")
        return None
    except json.JSONDecodeError:
        st.error(f"Error al decodificar la respuesta JSON para '{query}'")
        return None

def generar_definicion(termino, contexto):
    url = "https://api.together.xyz/inference"
    payload = json.dumps({
        "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
        "prompt": f"Contexto: {contexto}\n\nTérmino: {termino}\n\nProporciona una definición del término cultural '{termino}' según la visión de la cultura Maya. La definición debe ser más larga, detallada, e informativa, similar a una entrada de diccionario extendida. Incluye referencias a fuentes específicas que traten este concepto.\n\nDefinición:",
        "max_tokens": 2048,
        "temperature": 0.7,
        "top_p": 0.7,
        "top_k": 50,
        "repetition_penalty": 1,
        "stop": ["Término:"]
    })
    headers = {
        'Authorization': f'Bearer {TOGETHER_API_KEY}',
        'Content-Type': 'application/json'
    }
    response = requests.post(url, headers=headers, data=payload)
    return response.json()['output']['choices'][0]['text'].strip()

def create_docx(terminos_definiciones_fuentes):
    doc = Document()
    doc.add_heading('Diccionario Cultural - Visión Maya', 0)

    for termino, definicion, fuentes in terminos_definiciones_fuentes:
        doc.add_heading('Término', level=1)
        doc.add_paragraph(termino)

        doc.add_heading('Definición', level=2)
        doc.add_paragraph(definicion)

        # Evitar agregar una segunda instancia de "Fuentes"
        if fuentes:
            doc.add_heading('Fuentes', level=1)
            for fuente in fuentes:
                doc.add_paragraph(f"{fuente['author']}. ({fuente['year']}). *{fuente['title']}*. {fuente['journal']}, {fuente['volume']}({fuente['issue']}), {fuente['pages']}. {fuente['url']}", style='List Bullet')

    doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes académicas para un análisis más profundo.')

    return doc

def generar_todas_las_entradas():
    terminos_definiciones_fuentes = []

    for termino in terminos_culturales:
        # Buscar información relevante
        resultados_busqueda = buscar_informacion(termino)
        if resultados_busqueda is None:
            continue  # Saltar este término si hubo un error en la búsqueda
        contexto = "\n".join([item["snippet"] for item in resultados_busqueda.get("results", [])])
        fuentes = [{
            "author": item["author"] if "author" in item else "Autor desconocido",
            "year": item["year"] if "year" in item else "s.f.",
            "title": item["title"],
            "journal": item["journal"] if "journal" in item else "Revista desconocida",
            "volume": item["volume"] if "volume" in item else "",
            "issue": item["issue"] if "issue" in item else "",
            "pages": item["pages"] if "pages" in item else "",
            "url": item["url"]
        } for item in resultados_busqueda.get("results", [])]

        # Generar definición
        definicion = generar_definicion(termino, contexto)

        # Añadir a la lista de resultados
        terminos_definiciones_fuentes.append((termino, definicion, fuentes))

    # Crear y guardar el archivo DOCX
    doc = create_docx(terminos_definiciones_fuentes)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# UI para generación en batch
if st.button("Generar todas las entradas en batch"):
    with st.spinner("Generando todas las entradas del diccionario..."):
        doc_buffer = generar_todas_las_entradas()
        if doc_buffer:
            st.download_button(
                label="Descargar todas las definiciones en DOCX",
                data=doc_buffer,
                file_name="Diccionario_Cultural_Maya_Batch.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
