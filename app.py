import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO

# Set page configuration
st.set_page_config(page_title="Diccionario Cultural Maya", page_icon="📚", layout="wide")

# Function to create the information column
def crear_columna_info():
    st.markdown("""
    ## Sobre esta aplicación

    Esta aplicación es un Diccionario Cultural basado en la visión de la cultura Maya. Permite a los usuarios obtener definiciones de términos culturales según la interpretación de la cosmovisión maya.

    ### Cómo usar la aplicación:

    1. Elija un término cultural de la lista predefinida.
    2. Haga clic en "Generar entrada de diccionario" para obtener la definición desde la perspectiva maya.
    3. Lea la definición y las fuentes proporcionadas.
    4. Si lo desea, descargue un documento DOCX con toda la información.

    ### Autor y actualización:
    **Moris Polanco**, 28 ag 2024

    ### Cómo citar esta aplicación (formato APA):
    Polanco, M. (2024). *Diccionario Cultural Maya* [Aplicación web]. https://culturamaya.streamlit.app

    ---
    **Nota:** Esta aplicación utiliza inteligencia artificial para generar definiciones basadas en la visión maya. Verifique la información con fuentes adicionales para un análisis más profundo.
    """)

# Titles and Main Column
st.title("Diccionario Cultural Maya")

col1, col2 = st.columns([1, 2])

with col1:
    crear_columna_info()

with col2:
    TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
    SERPLY_API_KEY = st.secrets["SERPLY_API_KEY"]

    # 101 cultural terms related to the Maya perspective
    terminos_culturales = sorted([
         "Ajaw", "Balam (Jaguar)", "Ceiba (Árbol sagrado)", "Chaac (Dios de la lluvia)", "Ch'ulel (Espíritu o alma)",
        "Cosmovisión", "Creador", "Destino", "Dualidad", "Dzuli (Extranjero)", "Etnia", "Familia", "Hunab Ku (Dios supremo)",
        "Ik' (Viento)", "Itzamná (Dios del cielo)", "Ixchel (Diosa de la luna)", "Ja' (Agua)", "K'uh (Divinidad)",
        "Kukulcán (Serpiente emplumada)", "Lak'in (Punto cardinal este)", "Maíz (Elemento vital)", "Muerte",
        "Nahualismo (Creencia en el espíritu animal)", "Nación", "Naturaleza", "Noche", "Noj (Sabiduría)",
        "Oxlahun ti' k'uh (Trece dioses)", "Pawahtun (Dioses del inframundo)", "Persona", "Q'anil (Semilla)",
        "Religión", "Sak Nikte' (Flor blanca)", "Sak'ij (Iluminación)", "Sol", "Suerte", "Tezcatlipoca (Espejo humeante)",
        "Tierra", "Tikal (Ciudad maya)", "Toj (Tributo)", "Tzolkin (Calendario sagrado)", "Uinal (Mes maya)",
        "Uxmal (Tres veces construido)", "Viento", "Wak (Seis)", "Xibalbá (Inframundo maya)", "Yaxché (Árbol de la vida)",
        "Yuum k’aax (Dios de los bosques)", "Yuum ik’ (Dios del viento)", "Zama (Amanecer)", "Creación", "Espacio-tiempo",
        "Espíritu", "Fertilidad", "Fuego", "Identidad", "Inframundo", "Jaguar", "Kukulkan", "Luna", "Madera", "Maya",
        "Muerte", "Origen", "Paz", "Poder", "Pueblo", "Respiración", "Sacerdote", "Sociedad", "Sol", "Tiempo",
        "Tradición", "Trascendencia", "Unión", "Viento", "Vida", "Visión", "Voluntad", "Agua", "Animales", "Árboles",
        "Comunión", "Confianza", "Crecimiento", "Cultivo", "Fuerza", "Hombres", "Mujeres", "Naturaleza", "Puntos cardinales",
        "Reencarnación", "Ritual", "Sabiduría", "Salud", "Sueño", "Tierra", "Universo", "Vitalidad", "Xenil (Camino)",
        "Yuum Kaax (Dios de la flora y fauna)"
    ])

    def buscar_informacion(query):
        url = f"https://api.serply.io/v1/scholar/q={query}"
        headers = {
            'X-Api-Key': SERPLY_API_KEY,
            'Content-Type': 'application/json',
            'X-Proxy-Location': 'US',
            'X-User-Agent': 'Mozilla/5.0'
        }
        response = requests.get(url, headers=headers)
        return response.json()

    def generar_definicion(termino, contexto):
        url = "https://api.together.xyz/inference"
        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": f"Contexto: {contexto}\n\nTérmino: {termino}\n\nProporciona una definición del término cultural '{termino}' según la visión de la cultura Maya. La definición debe ser más larga, detallada, e informativa, similar a una entrada de diccionario extendida. Incluye referencias a fuentes específicas que traten este concepto.\n\nDefinición:",
            "max_tokens": 2048,
            "temperature": 0.7,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 1,
            "stop": ["Término:"]
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()['output']['choices'][0]['text'].strip()

    def create_docx(termino, definicion, fuentes):
        doc = Document()
        doc.add_heading('Diccionario Cultural - Visión Maya', 0)

        doc.add_heading('Término', level=1)
        doc.add_paragraph(termino)

        doc.add_heading('Definición', level=2)
        doc.add_paragraph(definicion)

        # Agregar "Fuentes" solo si hay fuentes disponibles
        if fuentes:
            doc.add_heading('Fuentes', level=1)
            for fuente in fuentes:
                doc.add_paragraph(f"{fuente['author']}. ({fuente['year']}). *{fuente['title']}*. {fuente['journal']}, {fuente['volume']}({fuente['issue']}), {fuente['pages']}. {fuente['url']}", style='List Bullet')

        doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes académicas para un análisis más profundo.')

        return doc

    st.write("Elige un término económico de la lista o propón tu propio término:")

    opcion = st.radio("", ["Elegir de la lista", "Proponer mi propio término"])

    if opcion == "Elegir de la lista":
        termino = st.selectbox("Selecciona un término:", terminos_culturales)
    else:
        termino = st.text_input("Ingresa tu propio término cultural:")

    if st.button("Generar entrada de diccionario"):
        if termino:
            with st.spinner("Buscando información y generando definición..."):
                # Buscar información relevante
                resultados_busqueda = buscar_informacion(termino)
                contexto = "\n".join([item["snippet"] for item in resultados_busqueda.get("results", [])])
                fuentes = [{
                    "author": item["author"] if "author" in item else "Autor desconocido",
                    "year": item["year"] if "year" in item else "s.f.",
                    "title": item["title"],
                    "journal": item["journal"] if "journal" in item else "Revista desconocida",
                    "volume": item["volume"] if "volume" in item else "",
                    "issue": item["issue"] if "issue" in item else "",
                    "pages": item["pages"] if "pages" in item else "",
                    "url": item["url"]
                } for item in resultados_busqueda.get("results", [])]

                # Generar definición
                definicion = generar_definicion(termino, contexto)

                # Mostrar la definición
                st.subheader(f"Definición para el término: {termino}")
                st.markdown(f"**{definicion}**")

                # Botón para descargar el documento
                doc = create_docx(termino, definicion, fuentes)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="Descargar definición en DOCX",
                    data=buffer,
                    file_name=f"Definicion_{termino.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("Por favor, selecciona un término.")
